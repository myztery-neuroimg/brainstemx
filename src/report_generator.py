"""
brainstemx.report_generator – GPT-4.1 vision → .docx

Requirements:
    pip install openai>=1.5 python-docx pandas pillow
"""

from __future__ import annotations
import base64, argparse, sys, logging
from datetime import date
from pathlib import Path
from typing import List, Dict, Tuple, Optional, Any
import pandas as pd, openai, docx
from docx.shared import Inches
import nibabel as nib, numpy as np, matplotlib.pyplot as plt

from .core import check_file_dependencies

MODEL = "gpt-4.1-vision-preview"

# ---------- helpers ----------------------------------------------------------
def to_b64(p:Path)->str:
    return "data:image/png;base64,"+base64.b64encode(p.read_bytes()).decode()

def sagittal(subj:Path)->Path:
    out=subj/"sag_overlay.png"
    if out.exists(): return out
    flair=nib.load(str(subj/"flair_to_t1.nii.gz")).get_fdata()
    mask =nib.load(str(subj/"lesion_sd2.0.nii.gz")).get_fdata()
    x=flair.shape[0]//2
    norm=(flair[x,:,:]-flair.min())/flair.ptp()
    rgba=np.stack([norm,norm,norm,np.ones_like(norm)],2)
    rgba[mask[x,:,:]>0]=[1,0,0,0.6]
    plt.imsave(out,(rgba*255).astype(np.uint8))
    return out

def collect_imgs(subj)->tuple[list[Path],list[dict]]:
    files=[subj/"QC_overlay.png", sagittal(subj)] + sorted(subj.glob("qc_*.png"))
    files=[p for p in files if p.exists()][:4]
    blocks=[{"type":"image_url","image_url":{"url":to_b64(p)}} for p in files]
    return files,blocks

def build_messages(csv_txt:str, blocks:list[dict])->list[dict]:
    system="You are a board-certified neuroradiologist. Write MRI brain-stem Findings & Impression."
    user=[{"type":"text","text":f"Lesion metrics CSV:\n```csv\n{csv_txt.strip()}\n```"}]+blocks
    return [{"role":"system","content":system},
            {"role":"user","content":user}]

def gpt_dictation(key:str,msgs)->str:
    openai.api_key=key
    res=openai.chat.completions.create(model=MODEL,messages=msgs,
                                       max_tokens=400,response_format={"type":"text"})
    return res.choices[0].message.content.strip()

def docx_report(subj:Path,text:str,imgs:list[Path])->Path:
    doc=docx.Document()
    doc.add_heading(f"BrainStemX auto-report – {subj.name}",1)
    doc.add_paragraph(f"Date: {date.today().isoformat()}")
    for line in text.splitlines(): doc.add_paragraph(line)
    doc.add_page_break(); doc.add_heading("Thumbnails",level=2)
    for p in imgs:
        doc.add_paragraph(p.name); doc.add_picture(str(p),width=Inches(4))
    fp=subj/f"{subj.name}_report.docx"; doc.save(fp); return fp

def check_prerequisites(subj_dir: Path) -> None:
    """Check that all required files exist before generating report.
    
    Args:
        subj_dir: Path to the subject's output directory
        
    Raises:
        FileNotFoundError: If any required files are missing
    """
    # Set up logging
    log = logging.getLogger("report_generator")
    if not log.handlers:
        log.setLevel(logging.INFO)
        handler = logging.StreamHandler()
        handler.setFormatter(logging.Formatter("%(levelname)s: %(message)s"))
        log.addHandler(handler)
    
    # Check for required files
    required_files = [
        subj_dir/"analysis.csv",
        subj_dir/"QC_overlay.png",
        subj_dir/"flair_to_t1.nii.gz",
        subj_dir/"lesion_sd2.0.nii.gz"
    ]
    
    try:
        check_file_dependencies(required_files, log, "report generation")
    except FileNotFoundError as e:
        raise FileNotFoundError(
            f"Cannot generate report: {str(e)}\n"
            f"Did you run the full pipeline on {subj_dir.name}?"
        ) from e

# ---------- main -------------------------------------------------------------
def generate(subj_dir: Path, api_key: str) -> Optional[Path]:
    """Generate a radiology report using GPT-4.1-vision model.
    
    Args:
        subj_dir: Path to the subject's output directory
        api_key: OpenAI API key
        
    Returns:
        Path to the generated DOCX report, or None if generation failed
        
    Raises:
        FileNotFoundError: If required files are missing
        ValueError: If API key is invalid
        Exception: For other errors during report generation
    """
    log = logging.getLogger("report_generator")
    
    try:
        # Verify prerequisites
        check_prerequisites(subj_dir)
        
        # Load data
        log.info(f"Generating report for {subj_dir.name}")
        csv_txt = (subj_dir/"analysis.csv").read_text()
        
        # Collect images and prepare blocks
        imgs, blocks = collect_imgs(subj_dir)
        if not imgs:
            log.warning("No images found for visualization")
            
        # Build messages and get AI response
        messages = build_messages(csv_txt, blocks)
        
        # Verify API key
        if not api_key or len(api_key) < 20:
            raise ValueError("Invalid OpenAI API key")
            
        # Get AI analysis
        log.info("Sending data to GPT-4.1-vision for analysis...")
        ai = gpt_dictation(api_key, messages)
        
        # Save outputs
        (subj_dir/"ai_report.txt").write_text(ai)
        fp = docx_report(subj_dir, ai, imgs)
        
        log.info(f"Report successfully generated: {fp}")
        return fp
        
    except FileNotFoundError as e:
        log.error(f"Missing files: {e}")
        raise
    except ValueError as e:
        log.error(f"Value error: {e}")
        raise
    except Exception as e:
        log.error(f"Error generating report: {e}")
        raise

def main():
    """Command-line entry point for report generation."""
    # Configure logging
    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
    
    # Parse arguments
    ap = argparse.ArgumentParser(description="BrainStem X Report Generator")
    ap.add_argument("--subj", required=True, help="Path to subject directory")
    ap.add_argument("--key", required=True, help="OpenAI API key")
    args = ap.parse_args()
    
    try:
        generate(Path(args.subj), args.key)
        return 0
    except Exception as e:
        print(f"ERROR: {e}", file=sys.stderr)
        return 1

if __name__ == "__main__":
    sys.exit(main())

